from __future__ import annotations

import io
import logging
from pathlib import Path

from server.schemas.config import ProviderConfig
from server.services.ocr_client import OcrError, is_image, parse_document

logger = logging.getLogger(__name__)

# A text-layer PDF shorter than this is treated as scanned (OCR candidate).
MIN_PDF_TEXT_CHARS = 100


class FileExtractError(ValueError):
    """Raised when a file cannot be turned into plain text."""


# Suffixes we treat as plain text and decode directly.
TEXT_SUFFIXES = {
    ".txt", ".md", ".markdown", ".json", ".csv", ".tsv", ".xml", ".html", ".htm",
    ".css", ".js", ".mjs", ".cjs", ".ts", ".jsx", ".tsx", ".py", ".java", ".kt",
    ".c", ".h", ".cpp", ".hpp", ".cs", ".go", ".rs", ".rb", ".php", ".swift",
    ".sql", ".yaml", ".yml", ".log", ".ini", ".conf", ".cfg", ".toml", ".sh",
    ".bat", ".ps1", ".properties", ".env", ".gradle", ".vue", ".svelte",
}


def extract_text(filename: str, raw: bytes, mime_type: str | None = None) -> str:
    """Extract plain text from an uploaded file's raw bytes.

    Text-like files are decoded directly. PDF and DOCX go through optional
    libraries that are imported lazily so the server still boots when they are
    not installed yet.
    """
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(raw)
    if suffix == ".docx":
        return _extract_docx(raw)
    if suffix in TEXT_SUFFIXES or (mime_type or "").startswith("text/"):
        return _decode_text(raw)

    # Unknown binary types: best-effort decode, never crash.
    return _decode_text(raw)


async def extract_text_async(
    filename: str,
    raw: bytes,
    mime_type: str | None = None,
    config: ProviderConfig | None = None,
) -> str:
    """Async extraction that can route PDFs/images through the OCR service.

    Behaviour:
    - OCR disabled (or no config): identical to ``extract_text``; images raise
      a clear error instead of producing garbage.
    - PDF + OCR ``auto``: only call OCR when pypdf yields little/no text
      (scanned document); keep whichever result is longer.
    - PDF + OCR ``always``: always try OCR, fall back to pypdf on failure.
    - Image + OCR: parse via OCR; surface a clear error if it fails.

    OCR failures never crash the upload: the plain-text result is used instead.
    """
    suffix = Path(filename).suffix.lower()
    ocr_enabled = bool(config and config.ocr_enabled)
    ocr_mode = getattr(config, "ocr_mode", "auto") if config else "auto"

    if suffix == ".pdf":
        base_text = _extract_pdf(raw)
        need_ocr = ocr_enabled and (
            ocr_mode == "always" or len(base_text.strip()) < MIN_PDF_TEXT_CHARS
        )
        if need_ocr:
            ocr_text = await _try_ocr(config, raw, filename)
            if ocr_text and len(ocr_text.strip()) >= len(base_text.strip()):
                return ocr_text
        return base_text

    if is_image(filename):
        if not ocr_enabled:
            raise FileExtractError("解析图片需要在设置中开启 OCR 文档解析")
        try:
            ocr_text = await parse_document(config, raw, filename)  # type: ignore[arg-type]
        except OcrError as exc:
            raise FileExtractError(f"图片 OCR 失败：{exc}") from exc
        if not ocr_text.strip():
            raise FileExtractError("OCR 未能从图片中识别出文本")
        return ocr_text

    return extract_text(filename, raw, mime_type)


async def _try_ocr(config: ProviderConfig | None, raw: bytes, filename: str) -> str | None:
    """Best-effort OCR that swallows errors so PDF extraction can fall back."""
    if config is None:
        return None
    try:
        return await parse_document(config, raw, filename)
    except OcrError as exc:
        logger.warning("OCR 解析失败，回退普通抽取：%s", exc)
        return None
    except Exception as exc:  # noqa: BLE001 - never let OCR break an upload
        logger.warning("OCR 解析异常，回退普通抽取：%s", exc)
        return None


def _decode_text(raw: bytes) -> str:
    for encoding in ("utf-8", "utf-8-sig", "gbk", "gb18030", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def _extract_pdf(raw: bytes) -> str:
    try:
        from pypdf import PdfReader
    except ImportError as exc:
        raise FileExtractError(
            "解析 PDF 需要 pypdf 依赖，请先执行 pip install -r server/requirements.txt"
        ) from exc

    try:
        reader = PdfReader(io.BytesIO(raw))
    except Exception as exc:  # noqa: BLE001 - pypdf raises a variety of errors
        raise FileExtractError(f"PDF 解析失败：{exc}") from exc

    parts: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            parts.append(text)
    return "\n\n".join(parts)


def _extract_docx(raw: bytes) -> str:
    try:
        import docx  # python-docx
    except ImportError as exc:
        raise FileExtractError(
            "解析 DOCX 需要 python-docx 依赖，请先执行 pip install -r server/requirements.txt"
        ) from exc

    try:
        document = docx.Document(io.BytesIO(raw))
    except Exception as exc:  # noqa: BLE001 - python-docx raises a variety of errors
        raise FileExtractError(f"DOCX 解析失败：{exc}") from exc

    lines: list[str] = [p.text for p in document.paragraphs if p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)
